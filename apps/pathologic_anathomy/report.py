from django.db.models import QuerySet
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def generate_report_header(report_name: str = ""):
    """Returns a document with the header of the hospital reports."""
    document = Document()
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.font.size = Pt(10)
    run.add_text("Hospital General Universitario")
    run.add_break()
    run.add_text("Vladimir Ilich Lenin")
    run.add_break()
    run = paragraph.add_run()
    run.font.size = Pt(12)
    run.add_text("Centro Oncológico Territorial de Holguín")
    run.add_break()
    run = paragraph.add_run()
    run.bold = True
    run.font.size = Pt(12)
    run.add_text(report_name)
    run.add_break()
    return document


def add_report_range(document, initial_date, final_date):
    """Add the date range to the report."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(
        f"Del {initial_date.strftime('%d/%m/%Y')} al {final_date.strftime('%d/%m/%Y')}"
    )
    run.italic = True
    run.font.size = Pt(12)


def add_data_table(document, data: QuerySet, columns, queryset_columns):
    """Adds the data table to the document."""
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.autofit = False
    row = table.rows[0].cells
    for index, text in enumerate(columns):
        row[index].text = text
        paragraph = row[index].paragraphs[0]
        paragraph.runs[0].font.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for obj in data:
        row = table.add_row().cells
        for index, column in enumerate(queryset_columns):
            row[index].text = str(getattr(obj, column))
            if len(columns) - 1 == index:
                paragraph = row[index].paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_total(document, data, column):
    """Adds the total to the document."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total = 0
    for obj in data:
        total += getattr(obj, column)
    run = paragraph.add_run(f"Total de casos: {total}")
    run.italic = True
    run.font.size = Pt(12)
    run.font.bold = True
